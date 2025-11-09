# skill_utils.py
"""
Skill utilities for SkillSense:
- cached model loading
- canonical skill loading & embeddings (cached)
- text extraction helpers for PDF/DOCX/TXT
- keyword matching and embedding-based fuzzy matching (numpy)
- GitHub profile parsing (public API)
- LinkedIn public profile minimal extraction (best-effort)
"""

from pathlib import Path
import json
import re
import numpy as np
import time
import streamlit as st

BASE_DIR = Path(__file__).parent

# -------------------------
# Cached model + canonical
# -------------------------
@st.cache_resource(show_spinner=False)
def get_sentence_transformer(model_name: str = "all-MiniLM-L6-v2"):
    # lazy import to avoid heavy startup until Streamlit session uses it
    from sentence_transformers import SentenceTransformer
    model = SentenceTransformer(model_name)
    return model

@st.cache_data(show_spinner=False)
def load_canonical_skills_and_embeddings(precomputed_emb_path: str = None, model_name: str = "all-MiniLM-L6-v2"):
    """
    Returns (skills_list, embeddings_numpy).
    If precomputed_emb_path provided and exists, loads .npy and skills_list.json.
    Otherwise computes embeddings with sentence-transformers and caches result.
    """
    skills_path = BASE_DIR / "skills_ontology.json"
    if not skills_path.exists():
        raise FileNotFoundError("skills_ontology.json not found in project root.")
    with open(skills_path, "r", encoding="utf8") as f:
        skills = json.load(f)

    # try loading precomputed files if provided
    if precomputed_emb_path:
        emb_path = Path(precomputed_emb_path)
        skills_list_path = emb_path.with_name("skills_list.json")
        if emb_path.exists() and skills_list_path.exists():
            try:
                embs = np.load(str(emb_path))
                # load skills list if present
                try:
                    skills_loaded = json.load(open(skills_list_path, "r", encoding="utf8"))
                    return skills_loaded, embs
                except Exception:
                    return skills, embs
            except Exception:
                pass

    # compute embeddings (slow first-time)
    model = get_sentence_transformer(model_name)
    embs = model.encode(skills, convert_to_numpy=True, show_progress_bar=False)
    return skills, embs

# -------------------------
# Text extraction helpers
# -------------------------
def extract_text_from_pdf_bytes(bytes_io):
    """
    bytes_io: an uploaded file-like object from Streamlit
    Returns string
    """
    import pdfplumber
    try:
        with pdfplumber.open(bytes_io) as pdf:
            text_pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(text_pages)
    except Exception as e:
        # fallback: try reading .read()
        try:
            raw = bytes_io.getvalue()
            return raw.decode("utf-8", errors="ignore")
        except Exception:
            raise e

def extract_text_from_docx(file_obj):
    """
    file_obj: uploaded file-like object from Streamlit
    Returns string
    """
    from docx import Document
    # Streamlit provides a SpooledTemporaryFile-like object; Document can accept file-like
    try:
        document = Document(file_obj)
        text = "\n".join([p.text for p in document.paragraphs])
        return text
    except Exception:
        # try saving to temp and reopening
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(file_obj.getvalue())
            tmp.flush()
            document = Document(tmp.name)
            return "\n".join([p.text for p in document.paragraphs])

def simple_text_cleanup(text: str):
    if not text:
        return ""
    text = text.lower()
    # keep typical punctuation for sentence split, remove odd chars
    text = re.sub(r'[^a-z0-9\.\,\-\n &]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# -------------------------
# Keyword & embedding match
# -------------------------
def keyword_match(text: str, canonical_skills=None):
    text_l = text.lower()
    found = set()
    if canonical_skills is None:
        canonical_skills, _ = load_canonical_skills_and_embeddings()
    for skill in canonical_skills:
        if skill and skill.lower() in text_l:
            found.add(skill)
    return sorted(found)

def cosine_sim_numpy(a: np.ndarray, b: np.ndarray):
    # a: (m, d), b: (n, d) -> (m, n)
    # normalize
    an = a / np.linalg.norm(a, axis=1, keepdims=True)
    bn = b / np.linalg.norm(b, axis=1, keepdims=True)
    return np.dot(an, bn.T)

def embedding_match(text: str, top_k: int = 6, threshold: float = 0.56, max_sentences: int = 40, model_name: str = "all-MiniLM-L6-v2", precomputed_emb_path: str = None):
    """
    Semantic fuzzy match using sentence-transformers embeddings and numpy cosine.
    - limits number of candidate sentences encoded to max_sentences (speed)
    - returns list of canonical skills sorted by score (descending)
    """
    # load canonical
    canonical_skills, canonical_embs = load_canonical_skills_and_embeddings(precomputed_emb_path, model_name)

    # split into candidate sentences/lines
    # give priority to lines with colons, dashes, bullets, parentheses, or that contain digits (years)
    lines = [l.strip() for l in re.split(r'[\n\r]', text) if l.strip()]
    # heuristics: keep lines that are likely to be useful
    scored_lines = []
    for ln in lines:
        score = 0
        if re.search(r'[-••–—:;]', ln):
            score += 1
        if re.search(r'\d{4}', ln):  # contains a year
            score += 1
        if len(ln.split()) <= 5:
            score += 1
        if re.search(r'(experience|developed|built|worked|managed|implemented|designed|deployed|using|utiliz)', ln.lower()):
            score += 2
        scored_lines.append((score, ln))
    # sort by heuristics descending and take top max_sentences
    scored_lines.sort(key=lambda x: -x[0])
    candidates = [ln for _, ln in scored_lines][:max_sentences]
    if not candidates:
        # fallback: split full text into sentences
        candidates = [s.strip() for s in re.split(r'[.\n]', text) if s.strip()][:max_sentences]
    # encode candidates (in batch)
    model = get_sentence_transformer(model_name)
    cand_embs = model.encode(candidates, convert_to_numpy=True, show_progress_bar=False)
    # compute cosine with canonical embs (numpy)
    sim = cosine_sim_numpy(cand_embs, canonical_embs)  # shape (len(candidates), len(skills))
    # for each canonical skill collect max similarity across candidates
    max_scores = np.max(sim, axis=0)  # shape (len(skills),)
    hits = {}
    for idx, score in enumerate(max_scores):
        if float(score) >= threshold:
            hits[canonical_skills[idx]] = float(score)
    # sort hits by score desc
    sorted_hits = sorted(hits.items(), key=lambda x: -x[1])
    return [s for s, sc in sorted_hits]

def extract_skills_from_text(text: str, use_embeddings: bool = True, threshold: float = 0.56, max_sentences: int = 40, precomputed_emb_path: str = None):
    """
    Combined exact keyword + embedding fuzzy matching.
    use_embeddings: toggle to speed up (fast mode: False)
    """
    text_clean = simple_text_cleanup(text)
    canonical_skills, _ = load_canonical_skills_and_embeddings(precomputed_emb_path)
    exact = set(keyword_match(text_clean, canonical_skills))
    embed = set()
    if use_embeddings:
        try:
            embed = set(embedding_match(text_clean, threshold=threshold, max_sentences=max_sentences, precomputed_emb_path=precomputed_emb_path))
        except Exception as e:
            # fallback to exact only
            embed = set()
    combined = sorted(list(exact.union(embed)))
    return combined

# -------------------------
# Compare helper
# -------------------------
def compare_user_with_role_skills(user_skills, role_skills):
    missing = list(set(role_skills) - set(user_skills))
    extra = list(set(user_skills) - set(role_skills))
    matching = list(set(user_skills).intersection(role_skills))
    score = len(matching) / len(role_skills) if role_skills else 0

    return {
        "missing": missing,
        "extra": extra,
        "matching": matching,
        "score": score
    }


# -------------------------
# GitHub & LinkedIn helpers
# -------------------------
def fetch_github_profile_readme(username: str, token: str = None, max_repos: int = 10):
    """
    Fetch the README or repo descriptions for a public GitHub username via GitHub API.
    Returns combined text of README and repo descriptions.
    token: optional personal access token to increase rate limit.
    """
    import requests
    headers = {}
    if token:
        headers["Authorization"] = f"token {token}"
    base = "https://api.github.com"
    # fetch user repos
    repos_url = f"{base}/users/{username}/repos?per_page={max_repos}&sort=updated"
    r = requests.get(repos_url, headers=headers, timeout=10)
    if r.status_code != 200:
        return ""
    items = r.json()
    pieces = []
    for repo in items:
        # repo description
        if repo.get("description"):
            pieces.append(repo.get("description"))
        # try README raw
        owner = repo.get("owner", {}).get("login")
        repo_name = repo.get("name")
        raw_url = f"https://raw.githubusercontent.com/{owner}/{repo_name}/master/README.md"
        try:
            rr = requests.get(raw_url, headers=headers, timeout=6)
            if rr.status_code == 200 and rr.text:
                pieces.append(rr.text[:5000])
        except Exception:
            pass
    return "\n\n".join(pieces)

def fetch_github_languages(username: str, token: str = None, max_repos: int = 20):
    """
    Returns aggregated language usage from user's public repos (dict language -> bytes)
    """
    import requests
    headers = {}
    if token:
        headers["Authorization"] = f"token {token}"
    base = "https://api.github.com"
    repos_url = f"{base}/users/{username}/repos?per_page={max_repos}&sort=updated"
    r = requests.get(repos_url, headers=headers, timeout=10)
    if r.status_code != 200:
        return {}
    repos = r.json()
    lang_agg = {}
    for repo in repos:
        langs_url = repo.get("languages_url")
        try:
            rr = requests.get(langs_url, headers=headers, timeout=6)
            if rr.status_code == 200:
                data = rr.json()
                for k, v in data.items():
                    lang_agg[k] = lang_agg.get(k, 0) + v
        except Exception:
            pass
    return lang_agg

def fetch_linkedin_public_text(profile_url: str, user_agent: str = None):
    """
    Best-effort public LinkedIn HTML fetch and visible text extraction.
    WARNING: LinkedIn actively blocks automated requests and requires login for full profiles.
    This function attempts to fetch and parse the public HTML. It may return empty or partial content.
    Prefer asking the user to paste their LinkedIn profile text or export as PDF.
    """
    import requests
    from bs4 import BeautifulSoup

    headers = {"User-Agent": user_agent or "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
    try:
        r = requests.get(profile_url, headers=headers, timeout=10)
        if r.status_code != 200:
            return ""
        soup = BeautifulSoup(r.text, "html.parser")
        # remove scripts/styles
        for script in soup(["script", "style", "noscript"]):
            script.extract()
        # join visible text
        text = soup.get_text(separator="\n")
        # do some cleanup
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        return "\n".join(lines[:500])  # limit to first 500 lines
    except Exception:
        return ""
